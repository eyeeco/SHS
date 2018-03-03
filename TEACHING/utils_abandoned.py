import docx
import time
import os.path as osp
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
from django.conf import settings


def export_homework(homework, student):
    document = Document()
    document.add_paragraph()
    document.add_paragraph()
    title = document.add_heading('人工智能大作业', 0)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()
    document.add_paragraph()
    document.add_paragraph()
    document.add_paragraph()
    course = document.add_paragraph()
    run = course.add_run('                                课  程  名：')
    run.font.size = Pt(18)
    run = course.add_run('          人工智能        ')
    run.font.size = Pt(18)
    run.font.underline = True
    institute = document.add_paragraph()
    run = institute.add_run('                                学 院（系）：')
    run.font.size = Pt(18)
    run = institute.add_run('          ' + str(
                            student.get_institute_display()) + '        ')
    run.font.size = Pt(18)
    run.font.underline = True
    education = document.add_paragraph()
    run = education.add_run('                                年       级：')
    run.font.size = Pt(18)
    run = education.add_run('          ' + str(
                            student.get_education_display()) + '        ')
    run.font.size = Pt(18)
    run.font.underline = True
    name = document.add_paragraph()
    run = name.add_run('                                学 生 姓 名：')
    run.font.size = Pt(18)
    run = name.add_run('          ' + str(student.user_info) + '        ')
    run.font.size = Pt(18)
    run.font.underline = True
    student_id = document.add_paragraph()
    run = student_id.add_run('                                学       号：')
    run.font.size = Pt(18)
    run = student_id.add_run('          ' + str(
                             student.student_id) + '        ')
    run.font.size = Pt(18)
    run.font.underline = True
    subtime = document.add_paragraph()
    run = subtime.add_run('                                时       间：')
    run.font.size = Pt(18)
    run = subtime.add_run('      ' + time.strftime('%Y-%m-%d', time.localtime(
                          time.time())) + '      ')
    run.font.size = Pt(18)
    run.font.underline = True
    document.add_paragraph()
    document.add_paragraph()
    document.add_paragraph()
    document.add_paragraph()
    document.add_paragraph()
    school = document.add_paragraph('大连理工大学')
    school.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    school_en = document.add_paragraph('Dalian University of Technology')
    school_en.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()
    for word in homework:
        file = docx.Document(word.file_field)
        for element in file.element.body:
            document.element.body.append(element)
    fname = '{}.docx'.format(student.user_info)
    save_path = osp.join(settings.TMP_FILES_ROOT, fname)
    document.save(save_path)
    url_path = osp.join(settings.TMP_FILES_URL, fname)
    return url_path


def export_allhomework(temp, student):
    document = Document()
    for stu in student:
        homework = temp[stu]
        for word in homework:
            file = docx.Document(word.file_field)
            for element in file.element.body:
                document.element.body.append(element)
    fname = 'all.docx'
    save_path = osp.join(settings.TMP_FILES_ROOT, fname)
    document.save(save_path)
    url_path = osp.join(settings.TMP_FILES_URL, fname)
    return url_path

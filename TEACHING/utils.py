import docx
import os.path as osp
from docx import Document
from django.conf import settings


def export_homework(homework, name):
    document = Document()
    for word in homework:
        file = docx.Document(word.file_field)
        for element in file.element.body:
            document.element.body.append(element)
    fname = '{}.docx'.format(name)
    save_path = osp.join(settings.TMP_FILES_ROOT, fname)
    document.save(save_path)
    url_path = osp.join(settings.TMP_FILES_URL, fname)
    return url_path


def export_allhomework(temp, student):
    document = Document()
    for stu in student:
        homework = temp[stu]
        for word in homework:
            file = docx.Document(word.file_field)
            for element in file.element.body:
                document.element.body.append(element)
    fname = 'all.docx'
    save_path = osp.join(settings.TMP_FILES_ROOT, fname)
    document.save(save_path)
    url_path = osp.join(settings.TMP_FILES_URL, fname)
    return url_path
